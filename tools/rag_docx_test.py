#!/usr/bin/env python3
"""RAG DOCX ingest — test recipe for the rig.

Like rag_pdf_test.py but for .docx. Extraction is fully verifiable and was
proven in the container (python-docx extracts paragraphs + tables, incl. Danish
letters). This confirms the full round trip on the rig where Ollama runs.

PREREQUISITES (on the rig):
    pip install python-docx
    # Ollama running with the embedding model:  ollama pull nomic-embed-text

RUN:
    python tools/rag_docx_test.py                  # generates + tests a sample
    python tools/rag_docx_test.py my_document.docx # tests your own .docx
"""
import base64, sys


def main() -> int:
    sys.path.insert(0, "worker")
    from app import rag_docx
    if not rag_docx.is_available():
        print("python-docx not installed. Run: pip install python-docx")
        return 1

    if len(sys.argv) > 1:
        with open(sys.argv[1], "rb") as f:
            data = f.read()
        print(f"using {sys.argv[1]}")
    else:
        import docx, io
        d = docx.Document()
        d.add_heading("ModelRig testdokument", 0)
        d.add_paragraph("Alva er Android-appen; motoren hedder ModelRig.")
        d.add_paragraph("Voice bruger faster-whisper til ASR og Piper til TTS.")
        t = d.add_table(rows=2, cols=2)
        t.rows[0].cells[0].text = "GPU"; t.rows[0].cells[1].text = "RTX 3060"
        t.rows[1].cells[0].text = "VRAM"; t.rows[1].cells[1].text = "12 GB"
        buf = io.BytesIO(); d.save(buf); data = buf.getvalue()
        print("using a generated Danish sample .docx (paragraphs + table)")

    ex = rag_docx.extract_text(data)
    print(f"\nEXTRACTED: {ex['paragraphs']} blocks, {ex['chars']} chars")
    print("--- text ---"); print(ex["text"][:400])

    import urllib.request, json
    b64 = base64.b64encode(data).decode()
    body = json.dumps({"docx_base64": b64, "source": "docx_test"}).encode()
    req = urllib.request.Request("http://127.0.0.1:8099/rag/ingest/docx",
                                 data=body, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=120) as r:
            result = json.loads(r.read())
        print(f"\nINGESTED: {result['chunks_added']} chunks, {result['total']} total")
    except Exception as e:
        print(f"\nINGEST failed: {e} (worker on :8099? Ollama on :11434?)")
        return 1

    q = "Hvilken GPU har rig-maskinen?"
    body = json.dumps({"query": q, "top_k": 3}).encode()
    req = urllib.request.Request("http://127.0.0.1:8099/rag/query",
                                 data=body, headers={"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=120) as r:
        ans = json.loads(r.read())
    print(f"\nQUERY: {q}\nANSWER: {ans.get('answer', ans)}")
    print("\nIf the answer mentions the RTX 3060, DOCX->RAG works (table content included).")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
