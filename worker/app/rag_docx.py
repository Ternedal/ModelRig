"""RAG DOCX ingest — text extraction with python-docx.

Mirrors rag_pdf.py exactly: extracts text from an uploaded .docx so it can go
through the EXISTING RAG ingest pipeline (chunk -> embed -> store). No new RAG
logic; just a text-extraction layer in front of rag.ingest().

python-docx reads Word's OOXML. We pull paragraph text AND table cell text
(tables often hold the actual data in real documents, and skipping them silently
loses content). Headings/paragraphs are joined with blank lines so the RAG
chunker still sees natural breaks.

OPTIONAL, like the PDF and Voice backends: python-docx is NOT a hard worker
dependency. Imported lazily; absent -> the DOCX endpoint returns a clean 501
with install instructions, and the rest of the worker is unaffected.

Note: only the modern .docx (OOXML) is supported, not the legacy binary .doc --
python-docx can't read .doc, so we detect and reject it with an honest message
rather than returning garbage. This module IS end-to-end testable without
special hardware (create a docx, extract, ingest, query).
"""
from __future__ import annotations

import logging

import io


def is_available() -> bool:
    """True if python-docx can be imported (installed)."""
    try:
        import docx  # noqa: F401  (python-docx)
        return True
    except Exception as exc:  # noqa: BLE001
        # Fail-closed is right: no python-docx, no DOCX-indlæsning. Silence is not.
        #
        # This swallows more than ImportError, and on the rig it will: a broken
        # python-docx wheel, a missing DLL, a CUDA mismatch. All of them arrive here
        # as "unavailable" with no reason, and the person reading the
        # capability list has to guess. F-501 was this exact shape -- an
        # `except Exception` hid an ImportError from a wrong class name for
        # eight releases, and the test passed because it asserted the failing
        # value and got it.
        logging.getLogger(__name__).info(
            "python-docx er ikke tilgængelig (DOCX-indlæsning slået fra): %r", exc)
        return False


def _iter_block_text(document) -> list[str]:
    """Collect paragraph text and table cell text from the document."""
    parts: list[str] = []
    for para in document.paragraphs:
        t = (para.text or "").strip()
        if t:
            parts.append(t)
    for table in document.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    return parts


def extract_text(docx_bytes: bytes) -> dict:
    """Extract text from .docx bytes.

    Returns {text, paragraphs, chars}. Raises RuntimeError (surfaced as 501/400
    by the endpoint) on a missing backend or an unreadable/legacy file.
    """
    try:
        import docx
    except Exception as e:
        raise RuntimeError(
            "DOCX ingest needs python-docx. Install it on the rig with: pip install python-docx"
        ) from e

    # Legacy binary .doc starts with the OLE magic (D0 CF 11 E0). python-docx
    # only reads OOXML .docx (a zip, starts with 'PK'). Be honest rather than
    # throwing an opaque zip error.
    if docx_bytes[:4] == b"\xd0\xcf\x11\xe0":
        raise RuntimeError("legacy .doc is not supported; save as .docx and retry")

    try:
        document = docx.Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise RuntimeError(f"could not open DOCX: {e}") from e

    parts = _iter_block_text(document)
    text = "\n\n".join(parts).strip()
    return {"text": text, "paragraphs": len(parts), "chars": len(text)}
